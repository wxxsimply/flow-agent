from docx import Document
from pathlib import Path

p = Path(r"d:/Code/flow-agent/docs")
docx = list(p.glob("*模板*.docx"))[0]
d = Document(str(docx))
print("FILE:", docx)
print("PARAS:", len(d.paragraphs), "TABLES:", len(d.tables))
for i, para in enumerate(d.paragraphs):
    t = para.text.strip()
    if t:
        print(f"P{i}: {t}")
for ti, table in enumerate(d.tables):
    print(f"TABLE{ti} rows={len(table.rows)} cols={len(table.columns)}")
    for ri, row in enumerate(table.rows):
        print(f"  R{ri}:", [c.text for c in row.cells])
