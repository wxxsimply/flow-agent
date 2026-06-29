"""Fill 伍孝轩 resume into 林炜豪 template — preserves icons, tables, layout."""

from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Twips
from docx.text.paragraph import Paragraph

PROJECT_BODY_INDENT = Twips(760)

from resume_content import ResumeData, default_resume

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
TEMPLATE_BLANK = DOCS / "林炜豪简历-模板-空白.docx"
OUT_DOCX = DOCS / "伍孝轩-Go后端简历.docx"
OUT_DOCX_NEW = DOCS / "伍孝轩-Go后端简历-new.docx"
OUT_PDF = DOCS / "伍孝轩-Go后端简历.pdf"

FONT_NAME = "微软雅黑"
BODY_SIZE = Pt(10.5)


def apply_run_font(run, *, bold: bool | None = None, size=None) -> None:
    run.font.name = FONT_NAME
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), FONT_NAME)
    if bold is not None:
        run.font.bold = bold
    if size is not None:
        run.font.size = size


def set_run_text_preserve_drawings(run, text: str) -> None:
    t_elems = run._element.findall(qn("w:t"))
    if t_elems:
        t_elems[0].text = text
        for elem in t_elems[1:]:
            elem.text = ""
    elif text:
        run.text = text


def set_paragraph_run_parts(paragraph: Paragraph, parts: list[str]) -> None:
    runs = paragraph.runs
    for idx, part in enumerate(parts):
        if idx < len(runs):
            set_run_text_preserve_drawings(runs[idx], part)
    for idx in range(len(parts), len(runs)):
        set_run_text_preserve_drawings(runs[idx], "")


def set_paragraph_text_keep_style(paragraph: Paragraph, text: str) -> None:
    bold = False
    size = None
    for run in paragraph.runs:
        if run.text.strip():
            bold = bool(run.font.bold)
            size = run.font.size
            break
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    apply_run_font(run, bold=bold, size=size or BODY_SIZE)


def set_bullet_line(paragraph: Paragraph, text: str) -> None:
    """Icon bullet paragraph (3 runs) — all visible text in first run."""
    runs = paragraph.runs
    if runs:
        set_run_text_preserve_drawings(runs[0], text)
        for run in runs[1:]:
            set_run_text_preserve_drawings(run, "")
    else:
        paragraph.add_run(text)


def set_award_line(paragraph: Paragraph, award: str) -> None:
    """Award lines use 3 runs (prefix / body / medal) — keep icons, replace text."""
    runs = paragraph.runs
    if len(runs) >= 3:
        if "金奖" in award:
            medal = "金奖"
        elif "银奖" in award:
            medal = "银奖"
        elif "铜奖" in award:
            medal = "铜奖"
        elif "一等奖" in award:
            medal = "一等奖"
        elif "三等奖" in award:
            medal = "三等奖"
        else:
            medal = runs[2].text or ""
        body = award
        for suffix in ("金奖", "银奖", "铜奖", "一等奖", "三等奖"):
            if award.endswith(suffix):
                body = award[: -len(suffix)]
                break
        prefix = ""
        for marker in ("ACM/CCPC", "ACM/ICPC", "蓝桥杯", "团体程序设计竞赛"):
            pos = body.find(marker)
            if pos != -1:
                prefix = body[: pos + len(marker)]
                body = body[pos + len(marker) :]
                break
        if not prefix:
            prefix = body[: min(12, len(body))]
            body = body[len(prefix) :]
        set_run_text_preserve_drawings(runs[0], prefix)
        set_run_text_preserve_drawings(runs[1], body)
        set_run_text_preserve_drawings(runs[2], medal)
    else:
        set_paragraph_run_parts(paragraph, [award])


def set_cell_text_preserve_drawings(cell, text: str) -> None:
    if not cell.paragraphs:
        return
    para = cell.paragraphs[0]
    if para.runs:
        set_run_text_preserve_drawings(para.runs[0], text)
    else:
        para.add_run(text)


def delete_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def clone_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_p = deepcopy(paragraph._element)
    paragraph._element.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def find_paragraph(doc: Document, prefix: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(prefix):
            return paragraph
    return None


def set_paragraph_indent(
    paragraph: Paragraph,
    *,
    left: Twips | None = None,
    first_line: Twips | None = None,
) -> None:
    if left is not None:
        paragraph.paragraph_format.left_indent = left
    if first_line is not None:
        paragraph.paragraph_format.first_line_indent = first_line


def clear_paragraph_indent(paragraph: Paragraph) -> None:
    paragraph.paragraph_format.left_indent = None
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.right_indent = None


def paragraph_has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._element.findall(".//" + qn("w:drawing")))


def remove_decorative_hlines_after_project(doc: Document) -> None:
    """Remove empty template divider lines between project body and skills."""
    anchor = find_paragraph(doc, "6.工程化交付")
    if anchor is None:
        return
    elem = anchor._element.getnext()
    while elem is not None and elem.tag.endswith("p"):
        para = Paragraph(elem, anchor._parent)
        if para.text.strip():
            break
        if paragraph_has_drawing(para):
            nxt = elem.getnext()
            elem.getparent().remove(elem)
            elem = nxt
            continue
        break


def fix_skills_section_layout(doc: Document) -> None:
    skills_table = None
    for table in doc.tables:
        if table.rows[0].cells[0].text.strip() == "专业技能":
            skills_table = table
            break
    if skills_table is None:
        return

    for cell in skills_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            clear_paragraph_indent(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    elem = skills_table._element.getnext()
    while elem is not None:
        if elem.tag.endswith("tbl"):
            break
        if not elem.tag.endswith("p"):
            elem = elem.getnext()
            continue
        paragraph = Paragraph(elem, skills_table._element.getparent())
        text = paragraph.text.strip()
        if text == "教育经历":
            break
        if text:
            clear_paragraph_indent(paragraph)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elem = elem.getnext()


def fix_exact_line_spacing(paragraph: Paragraph) -> None:
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is not None and spacing.get(qn("w:lineRule")) == "exact":
        spacing.set(qn("w:lineRule"), "auto")
        spacing.set(qn("w:line"), "240")


def normalize_document(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        fix_exact_line_spacing(paragraph)
        for run in paragraph.runs:
            apply_run_font(run, bold=run.font.bold, size=run.font.size or BODY_SIZE)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    fix_exact_line_spacing(paragraph)
                    for run in paragraph.runs:
                        apply_run_font(run, size=BODY_SIZE)


def fill_header(paragraph: Paragraph, data: ResumeData) -> None:
    parts = [
        data.name,
        "\n",
        f"手机号：{data.phone}|",
        f"邮箱：{data.email}",
        "",
        f" | 微信号：{data.wechat} ",
        f"\n{data.tagline}",
        "",
        "",
        "",
    ]
    set_paragraph_run_parts(paragraph, parts)


def fill_awards(doc: Document, awards: list[str]) -> None:
    anchor = find_paragraph(doc, "2025年")
    if anchor is None or not awards:
        raise RuntimeError("award slot not found")
    set_award_line(anchor, awards[0])
    insert_after = anchor
    next_elem = anchor._element.getnext()
    if next_elem is not None and next_elem.tag.endswith("p") and len(awards) > 1:
        second = Paragraph(next_elem, anchor._parent)
        set_award_line(second, awards[1])
        insert_after = second
        start = 2
    else:
        start = 1
    for award in awards[start:]:
        new_para = clone_paragraph_after(insert_after)
        set_award_line(new_para, award)
        insert_after = new_para


def remove_internship_section(doc: Document) -> None:
    prefixes = (
        "实习经历",
        "北京空间漫步",
        "公司致力于",
        "主要职责：参与",
        "1.服务架构",
        "、文件提交入库",
    )
    for paragraph in list(doc.paragraphs):
        if any(paragraph.text.strip().startswith(prefix) for prefix in prefixes):
            delete_paragraph(paragraph)


def remove_extra_projects(doc: Document) -> None:
    prefixes = (
        "设计面向短信",
        "5.实现定时清理",
        "基于当前的痛点",
        "2.CozeLoop",
        "4.集成MCP",
    )
    for paragraph in list(doc.paragraphs):
        if any(paragraph.text.strip().startswith(prefix) for prefix in prefixes):
            delete_paragraph(paragraph)
    for table in list(doc.tables):
        label = table.rows[0].cells[0].text.strip()
        if label == "专业技能":
            continue
        table._element.getparent().remove(table._element)


def clone_paragraph_after_element(template: Paragraph, after_element) -> Paragraph:
    new_p = deepcopy(template._element)
    after_element.addnext(new_p)
    return Paragraph(new_p, template._parent)


def fill_project_section(doc: Document, data: ResumeData) -> None:
    """Title + body paragraphs (no table); duty lines get left indent."""
    title_slot = find_paragraph(doc, "分布式容错KV")
    if title_slot is None:
        section = find_paragraph(doc, "项目经历")
        if section is None:
            raise RuntimeError("project section not found")
        title_slot = clone_paragraph_after(section)

    body_style = find_paragraph(doc, "2025年")
    if body_style is None:
        raise RuntimeError("template paragraph not found")

    set_paragraph_text_keep_style(
        title_slot,
        f"{data.project_title}\t{data.project_role}\t{data.project_period}",
    )
    if title_slot.paragraph_format.left_indent is None or title_slot.paragraph_format.left_indent == 0:
        set_paragraph_indent(title_slot, left=PROJECT_BODY_INDENT)

    insert_elem = title_slot._element
    blocks: list[tuple[str, bool]] = [
        (data.project_intro, True),
        (data.project_url, True),
        (data.project_label, True),
        *[(duty, True) for duty in data.project_duties],
    ]
    for text, indented in blocks:
        new_para = clone_paragraph_after_element(body_style, insert_elem)
        set_bullet_line(new_para, text)
        if indented:
            set_paragraph_indent(new_para, left=PROJECT_BODY_INDENT)
        insert_elem = new_para._element


def fill_skills_paragraph(doc: Document, skills: list[str]) -> None:
    slot = find_paragraph(doc, "1. 能够使用Gin")
    if slot is not None:
        delete_paragraph(slot)

    body_style = find_paragraph(doc, "分布式容错KV") or find_paragraph(doc, "2025年")
    if body_style is None:
        raise RuntimeError("skills body template not found")

    skills_table = None
    for table in doc.tables:
        if table.rows[0].cells[0].text.strip() == "专业技能":
            skills_table = table
            break
    if skills_table is None:
        raise RuntimeError("skills table not found")

    insert_elem = skills_table._element
    for skill in skills:
        new_para = clone_paragraph_after_element(body_style, insert_elem)
        set_bullet_line(new_para, skill)
        clear_paragraph_indent(new_para)
        new_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        insert_elem = new_para._element


def fill_education(doc: Document, data: ResumeData) -> None:
    school = find_paragraph(doc, "东莞理工学院")
    award_style = find_paragraph(doc, "2025年")
    if school and award_style:
        new_p = deepcopy(award_style._element)
        school._element.addprevious(new_p)
        set_bullet_line(Paragraph(new_p, school._parent), "教育经历")
    if school:
        set_paragraph_run_parts(school, [data.edu_school.strip()])
    major = find_paragraph(doc, "软件工程")
    if major:
        set_paragraph_run_parts(major, [data.edu_major])


def export_pdf(docx_path: Path, pdf_path: Path) -> None:
    import win32com.client

    for progid in ("Kwps.Application", "Word.Application"):
        try:
            app = win32com.client.Dispatch(progid)
            app.Visible = False
            try:
                doc = app.Documents.Open(str(docx_path.resolve()))
                doc.ExportAsFixedFormat(str(pdf_path.resolve()), 17)
                doc.Close(False)
            finally:
                app.Quit()
            return
        except Exception:
            continue
    raise RuntimeError("未找到 WPS/Word，无法导出 PDF")


def build(data: ResumeData | None = None, export_pdf_file: bool = False) -> Path:
    if not TEMPLATE_BLANK.exists():
        raise FileNotFoundError(f"template missing: {TEMPLATE_BLANK}")

    content = data or default_resume()
    out_path = OUT_DOCX_NEW
    try:
        shutil.copy2(TEMPLATE_BLANK, out_path)
    except PermissionError:
        out_path = DOCS / "伍孝轩-Go后端简历-new2.docx"
        shutil.copy2(TEMPLATE_BLANK, out_path)
    doc = Document(str(out_path))

    fill_header(doc.paragraphs[1], content)
    fill_awards(doc, content.awards)
    remove_internship_section(doc)
    remove_extra_projects(doc)
    fill_project_section(doc, content)
    remove_decorative_hlines_after_project(doc)
    fill_skills_paragraph(doc, content.skills)
    fix_skills_section_layout(doc)
    fill_education(doc, content)

    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip().startswith(("设计面向", "2.CozeLoop", "4.集成MCP", "5.实现定时")):
            delete_paragraph(paragraph)
    normalize_document(doc)

    try:
        doc.save(str(out_path))
    except PermissionError:
        alt = DOCS / "伍孝轩-Go后端简历-new2.docx"
        doc.save(str(alt))
        print(f"警告：{out_path.name} 被占用，已写入 {alt.name}")
        return alt

    if export_pdf_file:
        try:
            export_pdf(out_path, OUT_PDF)
        except Exception as exc:
            print(f"PDF 导出跳过：{exc}")

    return out_path


if __name__ == "__main__":
    path = build()
    from zipfile import ZipFile

    with ZipFile(path) as zf:
        media = [n for n in zf.namelist() if n.startswith("word/media/")]
    doc = Document(str(path))
    print(path)
    print(f"media={len(media)} tables={len(doc.tables)} paragraphs={len(doc.paragraphs)}")
