"""Reformat 伍孝轩 resume docx: same text, cleaner layout (v1.0.2-style)."""

from __future__ import annotations

import re
import shutil
from copy import deepcopy
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SRC = DOCS / "伍孝轩-Go后端简历.docx"
BACKUP = DOCS / "伍孝轩-Go后端简历.source.bak.docx"
TEMPLATE_BLANK = DOCS / "林炜豪简历-模板-空白.docx"
OUT = DOCS / "伍孝轩-Go后端简历.docx"

FONT_NAME = "微软雅黑"


@dataclass
class ParsedResume:
    name: str
    contact: str
    tagline: str
    awards: list[str] = field(default_factory=list)
    project_title: str = ""
    project_role: str = ""
    project_period: str = ""
    project_intro: str = ""
    project_url: str = ""
    project_label: str = "项目亮点："
    project_duties: list[str] = field(default_factory=list)
    skills: list[str] = field(default_factory=list)
    edu_school: str = ""
    edu_major: str = ""


def parse_source(doc: Document) -> ParsedResume:
    """Parse resume; supports both flat and already-tidied layouts."""
    blocks: list[tuple[str, list[str]]] = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        lines = [ln.strip() for ln in p.text.splitlines() if ln.strip()]
        blocks.append((text, lines))

    if not blocks:
        raise ValueError("empty resume")

    name = blocks[0][1][0]
    contact = blocks[0][1][1] if len(blocks[0][1]) > 1 else ""
    tagline = blocks[0][1][2] if len(blocks[0][1]) > 2 else ""

    idx = 1
    if idx < len(blocks) and "届" in blocks[idx][0] and not blocks[idx][0].startswith("20"):
        tagline = blocks[idx][0]
        idx += 1

    awards: list[str] = []
    while idx < len(blocks) and blocks[idx][0].startswith("20"):
        awards.append(blocks[idx][0])
        idx += 1

    project_title = project_role = project_period = ""
    project_intro = project_url = ""
    duties: list[str] = []

    if idx < len(blocks) and blocks[idx][0] == "项目经历":
        idx += 1
    if idx < len(blocks) and "FlowAgent" in blocks[idx][0]:
        header = blocks[idx][0]
        idx += 1
        period_m = re.search(r"(\d{4}年\d{1,2}月-\S+)", header)
        if period_m:
            project_period = period_m.group(1).replace(" ", "")
            header = header[: period_m.start()].strip()
        role = "独立开发"
        if header.endswith(role):
            project_title = header[: -len(role)].strip()
            project_role = role
        else:
            parts = header.split("\t")
            project_title = parts[0].strip() if parts else header
            project_role = parts[1].strip() if len(parts) > 1 else role
            if len(parts) > 2 and not project_period:
                project_period = parts[2].strip()

    if idx < len(blocks) and blocks[idx][0].startswith("该项目"):
        project_intro = blocks[idx][0]
        idx += 1
    if idx < len(blocks) and blocks[idx][0].startswith("项目地址"):
        project_url = blocks[idx][0]
        idx += 1
    if idx < len(blocks) and blocks[idx][0] == "主要工作：":
        idx += 1
    while idx < len(blocks) and re.match(r"^\d+\.", blocks[idx][0]):
        duties.append(blocks[idx][0])
        idx += 1

    skills: list[str] = []
    if idx < len(blocks) and blocks[idx][0] == "专业技能":
        idx += 1
    while idx < len(blocks):
        text, lines = blocks[idx]
        if text == "教育经历":
            idx += 1
            break
        if text.startswith("1.") and "Gin" in text:
            idx += 1
            continue
        for line in lines:
            if line == "教育经历":
                break
            if re.match(r"^\d+\.", line):
                skills.append(line)
        if skills:
            idx += 1
            break
        if re.match(r"^\d+\.", text):
            skills.append(text)
        idx += 1

    edu_school = edu_major = ""
    if idx < len(blocks) and "东莞理工" in blocks[idx][0]:
        edu_school = blocks[idx][0]
        idx += 1
    if idx < len(blocks):
        edu_major = blocks[idx][0]

    if not awards or not duties:
        raise ValueError("parsed resume incomplete — restore from backup")

    return ParsedResume(
        name=name,
        contact=contact,
        tagline=tagline,
        awards=awards,
        project_title=project_title,
        project_role=project_role,
        project_period=project_period,
        project_intro=project_intro,
        project_url=project_url,
        project_duties=duties,
        skills=skills,
        edu_school=edu_school,
        edu_major=edu_major,
    )


def canonical_content() -> ParsedResume:
    """Frozen copy of pre-tidy docx text (content unchanged)."""
    return ParsedResume(
        name="伍孝轩",
        contact="手机号：13727480563|邮箱：3110337564@qq.com | 微信号： wxx13798608302",
        tagline="28届本科|Golang|AI应用/AIGC工作流",
        awards=[
            "2025年ACM/CCPC福建省大学生程序设计竞赛全国邀请赛金奖",
            "2025年ACM/ICPC国际大学生程序设计竞赛武汉站全国邀请赛银奖",
            "2026年ACM/CCPC中国大学生程序设计竞赛广西站全国邀请赛银奖",
            "2025年ACM/CCPC中国大学生程序设计竞赛东北站全国邀请赛铜奖",
            "2025年ACM/CCPC广东省大学生程序设计竞赛邀请赛铜奖",
            "2026年ACM/CCPC广东省大学生程序设计竞赛邀请赛铜奖",
            "2025年蓝桥杯全国三等奖",
            "2026年团体程序设计竞赛团队全国三等奖",
            "2026年团体程序设计竞赛个人全国三等奖",
            "2025年团体程序设计竞赛团队全国三等奖",
        ],
        project_title="FlowAgent（AI微电影/短视频全自动生产Agent）",
        project_role="独立开发",
        project_period="2025年05月-至今",
        project_intro=(
            "该项目是基于 Go 自研的全链路 AIGC Agent，面向 网文连载 → 分镜 → AI 音画合成 → 合规发布 场景，"
            "将扩写、分镜、制片、合规、发布拆为可恢复、可验收的阶段流水线；支持 CLI 与 Desktop 双入口，"
            "产物落盘 runs/<run_id>/artifacts/，可通过 flowagent resume --from-stage 断点续跑。"
            "基于 YAML 阶段机编排七段流水线，产物可验收、可断点续跑。"
        ),
        project_url="项目地址：flow-agent（本地 monorepo 项目）",
        project_label="项目亮点：",
        project_duties=[
            "1.工作流引擎：YAML 七阶段编排，阶段门禁与多 Provider 切换。",
            "2.分镜 Agent：首镜 JSON 扩写，Skills 注入与分镜自动审查修复。",
            "3.多模态制片：TTS+文生图+i2v+FFmpeg 合成，Provider 故障自动降级。",
            "4.视频质量优化：physics 约束与 PropLock，修复道具穿模与跨镜一致性。",
            "5.连载记忆与成本：SQLite FTS 知识库与 cost-ledger 用量统计归档。",
            "6.工程化交付：CLI/Desktop 双入口，核心模块单元测试覆盖。",
        ],
        skills=[
            "1. 熟练 Go 后端与 Agent 开发，掌握 goroutine、errgroup 模块化设计。",
            "2. 熟悉 LLM 工程化：Prompt 调优、结构化 JSON 与 Skill 注入。",
            "3. 熟悉 AIGC 多模态：DashScope/Seedance，TTS、文生图与图生视频。",
            "4. 熟悉 YAML 工作流、状态机、Provider 插件化与 SQLite 全文检索。",
            "5. 熟悉并发编程：channel、条件变量与协程。",
            "6. 熟悉 Linux/Git，可独立部署服务器与协同开发。",
            "7. 熟悉数据结构算法：树、图、搜索、DP 等，有竞赛基础。",
            "8. 熟悉计算机网络与操作系统。",
            "9. 热爱新技术，抗压强，擅长 AI Agent/MCP/Skill 辅助开发。",
        ],
        edu_school="东莞理工学院（一本）\t2024年9月-2028年6月",
        edu_major="大数据技术与应用",
    )


def load_content() -> ParsedResume:
    if BACKUP.exists():
        try:
            return parse_source(Document(str(BACKUP)))
        except ValueError:
            pass
    if SRC.exists():
        try:
            data = parse_source(Document(str(SRC)))
            if not BACKUP.exists():
                shutil.copy2(SRC, BACKUP)
            return data
        except ValueError:
            pass
    return canonical_content()


def delete_paragraph(paragraph: Paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def set_paragraph_text_keep_style(paragraph: Paragraph, text: str) -> None:
    bold = False
    size = None
    for run in paragraph.runs:
        if run.text.strip():
            bold = bool(run.font.bold)
            size = run.font.size
            break
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    run = paragraph.add_run(text)
    apply_run_font(run, bold=bold, size=size)


def apply_run_font(run, *, bold: bool | None = None, size=None) -> None:
    run.font.name = FONT_NAME
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), FONT_NAME)
    if bold is not None:
        run.font.bold = bold
    if size is not None:
        run.font.size = size


def clone_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = deepcopy(paragraph._element)
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    set_paragraph_text_keep_style(new_para, text)
    return new_para


def find_paragraph(doc: Document, prefix: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def fix_exact_line_spacing(paragraph: Paragraph) -> None:
    p_pr = paragraph._element.find(qn("w:pPr"))
    if p_pr is None:
        return
    spacing = p_pr.find(qn("w:spacing"))
    if spacing is not None and spacing.get(qn("w:lineRule")) == "exact":
        spacing.set(qn("w:lineRule"), "auto")
        spacing.set(qn("w:line"), "240")


def normalize_document(doc: Document) -> None:
    for p in doc.paragraphs:
        fix_exact_line_spacing(p)
        for run in p.runs:
            apply_run_font(run, bold=run.font.bold, size=run.font.size)


def remove_internship_and_extra(doc: Document) -> None:
    prefixes = (
        "实习经历",
        "北京空间漫步",
        "公司致力于",
        "主要职责：参与",
        "1.服务架构",
        "、文件提交入库",
        "设计面向短信",
        "5.实现定时清理",
        "基于当前的痛点",
        "2.CozeLoop",
        "4.集成MCP",
        "1. 能够使用Gin",
    )
    for p in list(doc.paragraphs):
        if any(p.text.strip().startswith(x) for x in prefixes):
            delete_paragraph(p)
    for table in list(doc.tables):
        label = table.rows[0].cells[0].text.strip()
        if label == "专业技能":
            continue
        table._element.getparent().remove(table._element)


def remove_empty_paragraphs(doc: Document) -> None:
    for p in list(doc.paragraphs):
        if not p.text.strip():
            delete_paragraph(p)


def fill_header(doc: Document, data: ParsedResume) -> None:
    header = data.name
    if data.contact:
        header += f"\n{data.contact}"
    if data.tagline:
        header += f"\n{data.tagline}"
    set_paragraph_text_keep_style(doc.paragraphs[1], header)


def fill_awards(doc: Document, awards: list[str]) -> None:
    anchor = find_paragraph(doc, "2025年")
    if anchor is None or not awards:
        raise RuntimeError("award slot not found")
    set_paragraph_text_keep_style(anchor, awards[0])
    insert_after = anchor
    nxt = anchor._element.getnext()
    if nxt is not None and nxt.tag.endswith("p") and len(awards) > 1:
        para2 = Paragraph(nxt, anchor._parent)
        set_paragraph_text_keep_style(para2, awards[1])
        insert_after = para2
        start = 2
    else:
        start = 1
    for award in awards[start:]:
        insert_after = clone_paragraph_after(insert_after, award)


def fill_project(doc: Document, data: ParsedResume) -> None:
    slot = find_paragraph(doc, "分布式容错KV")
    if slot is None:
        title = find_paragraph(doc, "项目经历")
        if title is None:
            raise RuntimeError("project slot not found")
        slot = clone_paragraph_after(title, "")
    set_paragraph_text_keep_style(
        slot,
        f"{data.project_title}\t{data.project_role}\t{data.project_period}",
    )
    insert_after = slot
    if data.project_intro:
        insert_after = clone_paragraph_after(insert_after, data.project_intro)
    if data.project_url:
        insert_after = clone_paragraph_after(insert_after, data.project_url)
    if data.project_duties:
        label = data.project_label or "项目亮点："
        insert_after = clone_paragraph_after(insert_after, label)
        for duty in data.project_duties:
            insert_after = clone_paragraph_after(insert_after, duty)


def fill_skills(doc: Document, skills: list[str]) -> None:
    if not skills:
        return
    anchor = find_paragraph(doc, "6.工程化交付") or find_paragraph(doc, "5.连载记忆")
    if anchor is None:
        anchor = find_paragraph(doc, "主要工作：")
    if anchor is None:
        raise RuntimeError("skills anchor not found")
    title = clone_paragraph_after(anchor, "专业技能")
    insert_after = clone_paragraph_after(title, skills[0])
    for item in skills[1:]:
        insert_after = clone_paragraph_after(insert_after, item)
    clone_paragraph_after(insert_after, "教育经历")


def fill_education(doc: Document, school: str, major: str) -> None:
    edu_school = find_paragraph(doc, "东莞理工学院")
    if edu_school and school:
        set_paragraph_text_keep_style(edu_school, school)
    edu_major = find_paragraph(doc, "软件工程")
    if edu_major and major:
        set_paragraph_text_keep_style(edu_major, major)


def ensure_project_section_title(doc: Document, awards: list[str]) -> None:
    if find_paragraph(doc, "项目经历"):
        return
    anchor = find_paragraph(doc, awards[-1]) if awards else None
    if anchor:
        clone_paragraph_after(anchor, "项目经历")


def build() -> Path:
    data = load_content()
    doc = Document(str(TEMPLATE_BLANK))

    fill_header(doc, data)
    fill_awards(doc, data.awards)
    remove_internship_and_extra(doc)
    ensure_project_section_title(doc, data.awards)
    fill_project(doc, data)
    fill_skills(doc, data.skills)
    fill_education(doc, data.edu_school, data.edu_major)
    remove_empty_paragraphs(doc)
    normalize_document(doc)

    try:
        doc.save(str(OUT))
    except PermissionError:
        alt = DOCS / "伍孝轩-Go后端简历-tidy.docx"
        doc.save(str(alt))
        print(f"警告：{OUT.name} 被占用，已写入 {alt.name}")
        return alt
    return OUT


if __name__ == "__main__":
    print(build())
