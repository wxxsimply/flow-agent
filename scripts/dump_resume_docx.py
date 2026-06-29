"""Dump resume docx paragraphs to UTF-8 text for inspection."""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "伍孝轩-Go后端简历.docx"
OUT = ROOT / "docs" / "_resume_dump.txt"


def main() -> None:
    doc = Document(str(DOCX))
    lines: list[str] = []
    lines.append(f"PARAS={len(doc.paragraphs)} TABLES={len(doc.tables)}")
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        if not t.strip():
            lines.append(f"P{i}: (empty)")
            continue
        lines.append(f"P{i} [{len(t)} chars]:")
        for j, line in enumerate(t.splitlines()):
            lines.append(f"  L{j}: {line}")
    for ti, table in enumerate(doc.tables):
        lines.append(f"TABLE{ti}:")
        for ri, row in enumerate(table.rows):
            cells = [c.text.replace("\n", " | ") for c in row.cells]
            lines.append(f"  R{ri}: {cells}")
    OUT.write_text("\n".join(lines), encoding="utf-8")
    print(OUT)


if __name__ == "__main__":
    main()
